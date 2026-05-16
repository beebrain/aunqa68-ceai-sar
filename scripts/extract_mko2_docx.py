from __future__ import annotations

from pathlib import Path
import re

from docx import Document


def norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "aun-qa-wiki" / "raw" / "mko2-extract.md"
    out.parent.mkdir(parents=True, exist_ok=True)

    # Prefer a stable ASCII filename to avoid Windows encoding issues.
    # If not present, fall back to any DOCX containing "Comment".
    src = root / "mko2-comment.docx"
    if not src.exists():
        candidates = [p for p in root.glob("*.docx") if re.search(r"Comment", p.name, re.IGNORECASE)]
        if not candidates:
            raise SystemExit("No DOCX with 'Comment' found in project root.")
        src = candidates[0]

    doc = Document(str(src))

    lines: list[str] = []
    lines.append(f"# Extract from: {src.name}")
    lines.append("")

    lines.append("## Paragraphs")
    lines.append("")
    for i, p in enumerate(doc.paragraphs, 1):
        text = norm_ws(p.text)
        if not text:
            continue
        lines.append(f"- P{i}: {text}")

    lines.append("")
    lines.append("## Tables")
    lines.append("")
    for ti, table in enumerate(doc.tables, 1):
        lines.append(f"### Table {ti}")
        for ri, row in enumerate(table.rows, 1):
            cells = [norm_ws(c.text) for c in row.cells]
            lines.append(f"- R{ri}: " + " | ".join(cells))
        lines.append("")

    out.write_text("\n".join(lines), encoding="utf-8")
    print(f"WROTE {out}")


if __name__ == "__main__":
    main()

