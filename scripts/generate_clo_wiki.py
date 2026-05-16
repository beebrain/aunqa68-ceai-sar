"""Generate complete CLO wiki page from รวมเล่ม CLO รายวิชา.docx"""
from __future__ import annotations
from docx import Document
from pathlib import Path
import re

root = Path(__file__).resolve().parents[1]
src = root / "clo-courses-temp.docx"


def norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


# ---------- Load DOCX ----------
doc = Document(str(src))

paras = []
for j, p in enumerate(doc.paragraphs):
    text = norm_ws(p.text)
    if text:
        paras.append((j, text, p.style.name))

tables = []
for table in doc.tables:
    rows = [[norm_ws(c.text) for c in row.cells] for row in table.rows]
    tables.append(rows)

# ---------- Identify table types ----------
clo_idx: list[int] = []
bloom_idx: list[int] = []
act_idx: list[int] = []

for i, rows in enumerate(tables):
    if not rows:
        continue
    r0 = " ".join(rows[0])
    if rows[0][0].strip().startswith("CLO1") and len(rows[0]) == 2:
        clo_idx.append(i)
    elif "Object ของ Verb" in r0 or ("Learning Domain" in r0 and "Action Verb" in r0):
        bloom_idx.append(i)
    elif "Learning Environment" in r0 or "Leaning Environment" in r0:
        act_idx.append(i)

assert len(clo_idx) == 33, f"Expected 33 CLO tables, got {len(clo_idx)}"
assert len(bloom_idx) == 33, f"Expected 33 Bloom tables, got {len(bloom_idx)}"

# ---------- Extract course info from paragraphs ----------
course_sections: list[dict] = []
i = 0
while i < len(paras):
    _, text, _ = paras[i]
    if re.match(r"^1\. รหัสวิชา", text):
        code = text.replace("1. รหัสวิชา", "").strip()
        name = credits = group = desc = ""
        for j in range(i + 1, min(i + 25, len(paras))):
            t = paras[j][1]
            if t.startswith("2. ชื่อวิชา"):
                name = t.replace("2. ชื่อวิชา", "").strip()
            elif t.startswith("3. จำนวนหน่วยกิต"):
                credits = t.replace("3. จำนวนหน่วยกิต", "").strip()
            elif t.startswith("4. กลุ่มวิชา"):
                group = t.replace("4. กลุ่มวิชา", "").strip()
            elif t.startswith("6. คำอธิบายรายวิชา"):
                for k in range(j + 1, min(j + 5, len(paras))):
                    tt = paras[k][1]
                    if tt and not tt.startswith("7.") and not re.match(r"^[A-Z][a-z].*\.", tt):
                        desc = tt
                        break
                break
        course_sections.append({"code": code, "name": name, "credits": credits, "group": group, "desc": desc})
    i += 1

assert len(course_sections) == 33, f"Expected 33 courses, got {len(course_sections)}"


# ---------- Parse CLO list table ----------
def parse_clo_table(rows: list[list[str]]) -> list[dict]:
    clos = []
    for row in rows:
        if len(row) >= 2 and re.match(r"^CLO\d+", row[0]):
            txt = row[1]
            txt = re.sub(r"\s*\(verb\)", "", txt)
            txt = re.sub(r"\s*\(object\)", "", txt)
            txt = re.sub(r"\s*\(ส่วนขยาย\)", "", txt)
            txt = re.sub(r"\s*\(qualifier\)", "", txt)
            txt = norm_ws(txt)
            clos.append({"id": row[0].strip(), "text": txt})
    return clos


# ---------- Parse Bloom's table ----------
def parse_bloom_table(rows: list[list[str]]) -> list[dict]:
    start = 0
    for j, row in enumerate(rows):
        if row and ("Object ของ Verb" in row[0] or row[0] == "Object"):
            start = j + 1
    blooms = []
    for row in rows[start:]:
        if len(row) >= 4 and row[0] and row[0] not in ("Object ของ Verb", "Object"):
            blooms.append({"object": row[0], "domain": row[1], "verb": row[2], "qualifier": row[3]})
    return blooms


# ---------- Build CLO → Bloom + PLO mapping ----------
def build_clo_mapping(blooms: list[dict]) -> tuple[dict, dict]:
    clo_plo: dict[str, list[str]] = {}
    clo_bloom: dict[str, str] = {}
    for b in blooms:
        q = b["qualifier"]
        domain = b["domain"]
        # Extract Bloom level from domain string
        dm = re.search(r"^([A-Za-z]+)(?:\s*:|:)", domain)
        bloom_lv = dm.group(1) if dm else domain[:3]

        # Find [CLO#] references
        clo_refs = re.findall(r"CLO(\d+)", q)
        # Find PLO references
        plo_refs = re.findall(r"PLO\d+", q)

        for cr in clo_refs:
            key = f"CLO{cr}"
            if key not in clo_bloom:
                clo_bloom[key] = bloom_lv
            if key not in clo_plo:
                clo_plo[key] = []
            for p in plo_refs:
                if p not in clo_plo[key]:
                    clo_plo[key].append(p)
    return clo_bloom, clo_plo


# ---------- Render one course ----------
def render_course(cs: dict, clos: list[dict], blooms: list[dict]) -> list[str]:
    lines: list[str] = []
    lines.append(f"### {cs['code']} {cs['name']}")
    lines.append("")
    lines.append(f"- **หน่วยกิต**: {cs['credits']}")
    if cs.get("group"):
        lines.append(f"- **กลุ่มวิชา**: {cs['group']}")
    lines.append("")
    if cs.get("desc"):
        lines.append("#### คำอธิบายรายวิชา")
        lines.append("")
        lines.append(cs["desc"])
        lines.append("")

    clo_bloom, clo_plo = build_clo_mapping(blooms)

    if clos:
        lines.append("#### CLOs (รายข้อ)")
        lines.append("")
        lines.append("| CLO | คำอธิบาย | Bloom's | PLO |")
        lines.append("|-----|----------|---------|-----|")
        for clo in clos:
            cn = clo["id"]
            bloom_lv = clo_bloom.get(cn, "–")
            plos = clo_plo.get(cn, [])
            plo_str = ", ".join(plos) if plos else "–"
            lines.append(f"| {cn} | {clo['text']} | {bloom_lv} | {plo_str} |")
        lines.append("")

    lines.append("---")
    lines.append("")
    return lines


# ---------- Generate output ----------
MANDATORY = {"7015101", "7015102", "7015906", "7015103", "7015907", "7015908"}
THESIS = {"7015901", "7015902", "7015903", "7015904", "7015905"}

sections: dict[str, list] = {"mandatory": [], "thesis": [], "elective": []}
for n, cs in enumerate(course_sections):
    code = cs["code"]
    clos = parse_clo_table(tables[clo_idx[n]])
    blooms = parse_bloom_table(tables[bloom_idx[n]])
    entry = {"cs": cs, "clos": clos, "blooms": blooms}
    if code in MANDATORY:
        sections["mandatory"].append(entry)
    elif code in THESIS:
        sections["thesis"].append(entry)
    else:
        sections["elective"].append(entry)

out_lines: list[str] = [
    "---",
    "title: CLO รายวิชา (แสดงทุกรายวิชา แบบละเอียด)",
    "created: 2026-05-07",
    "updated: 2026-05-11",
    "type: summary",
    "tags: [curriculum, clo, course, mapping]",
    "sources: [raw/clo-all-courses-extract.md]",
    "confidence: high",
    "contested: false",
    "---",
    "",
    "# CLO รายวิชา (แสดงทุกรายวิชา แบบละเอียด)",
    "",
    "> ข้อมูล CLO ครบทุกรายวิชา สกัดจาก `รวมเล่ม CLO รายวิชา.docx` "
    "หลักสูตรวิศวกรรมศาสตรมหาบัณฑิต สาขาวิศวกรรมคอมพิวเตอร์และปัญญาประดิษฐ์ พ.ศ. 2568",
    "",
    "## 1) วิชาเฉพาะด้านบังคับ",
    "",
]
for e in sections["mandatory"]:
    out_lines.extend(render_course(e["cs"], e["clos"], e["blooms"]))

out_lines += ["## 2) วิทยานิพนธ์และสารนิพนธ์", ""]
for e in sections["thesis"]:
    out_lines.extend(render_course(e["cs"], e["clos"], e["blooms"]))

out_lines += ["## 3) วิชาเฉพาะด้านเลือก", ""]
for e in sections["elective"]:
    out_lines.extend(render_course(e["cs"], e["clos"], e["blooms"]))

out = root / "aun-qa-wiki" / "summaries" / "clo-all-courses-detailed.md"
out.write_text("\n".join(out_lines), encoding="utf-8")
print(f"WROTE {out}")
print(f"Lines: {len(out_lines)}")
print(f"Mandatory: {len(sections['mandatory'])} | Thesis: {len(sections['thesis'])} | Elective: {len(sections['elective'])}")
