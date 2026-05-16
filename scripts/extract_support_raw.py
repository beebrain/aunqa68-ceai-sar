"""Extract raw content from support data DOCX for inspection."""
from __future__ import annotations
from pathlib import Path
import docx

root = Path(__file__).resolve().parents[1]
src = root / "support-data-temp.docx"

doc = docx.Document(str(src))

lines: list[str] = ["# Raw Extraction: ข้อมูลพื้นฐาน AUN-QA สายสนับสนุน", ""]

# Paragraphs
lines.append("## PARAGRAPHS")
lines.append("")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        lines.append(f"[P{i:04d}] [{p.style.name}] {txt}")

lines.append("")
lines.append(f"Total paragraphs with text: {sum(1 for p in doc.paragraphs if p.text.strip())}")
lines.append("")

# Tables
lines.append("## TABLES")
lines.append("")
for ti, table in enumerate(doc.tables):
    lines.append(f"### Table {ti}")
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        # deduplicate merged cells
        seen = []
        prev = None
        for c in cells:
            if c != prev:
                seen.append(c)
            prev = c
        lines.append("| " + " | ".join(seen) + " |")
    lines.append("")

lines.append(f"Total tables: {len(doc.tables)}")

out = root / "aun-qa-wiki" / "raw" / "support-data-extract.md"
out.parent.mkdir(parents=True, exist_ok=True)
out.write_text("\n".join(lines), encoding="utf-8")
print(f"WROTE {out}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
